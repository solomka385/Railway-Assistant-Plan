"""
Скрипт для конвертации Markdown файла в формат Word (.docx)
Использование: python md_to_word.py <input_md_file> <output_docx_file>
"""

import re
import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def parse_markdown_to_docx(md_file_path, docx_file_path):
    """
    Конвертирует Markdown файл в Word документ с форматированием.
    """
    # Читаем Markdown файл
    with open(md_file_path, 'r', encoding='utf-8') as f:
        md_content = f.read()

    # Создаем Word документ
    doc = Document()

    # Устанавливаем шрифт по умолчанию
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Разбиваем содержимое на строки
    lines = md_content.split('\n')

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # Пустая строка
        if not line:
            doc.add_paragraph()
            i += 1
            continue

        # Заголовки
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            if level <= 6:
                heading_text = line.lstrip('#').strip()
                heading = doc.add_heading(heading_text, level=level)
                # Форматирование заголовка
                for run in heading.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(16 - level)
                    run.font.bold = True
                i += 1
                continue

        # Горизонтальная линия
        if line.strip() == '---':
            doc.add_paragraph('_' * 50)
            i += 1
            continue

        # Таблица (Markdown формат)
        if line.startswith('|') and line.endswith('|'):
            # Собираем все строки таблицы
            table_lines = []
            while i < len(lines) and lines[i].startswith('|'):
                table_lines.append(lines[i])
                i += 1

            # Пропускаем строку разделителя (|---|---|)
            if len(table_lines) > 1 and '---' in table_lines[1]:
                # Парсим таблицу
                rows = []
                for table_line in table_lines:
                    if '---' in table_line:
                        continue
                    # Убираем начальный и конечный |
                    cells = table_line[1:-1].split('|')
                    rows.append([cell.strip() for cell in cells])

                if rows:
                    # Создаем таблицу в Word
                    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                    table.style = 'Table Grid'

                    for row_idx, row_data in enumerate(rows):
                        for col_idx, cell_data in enumerate(row_data):
                            cell = table.rows[row_idx].cells[col_idx]
                            cell.text = cell_data
                            # Выравнивание текста в ячейке
                            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                            # Шрифт в ячейке
                            for run in cell.paragraphs[0].runs:
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(11)

                    # Добавляем отступ после таблицы
                    doc.add_paragraph()
                continue

        # Кодовый блок
        if line.startswith('```'):
            lang = line[3:].strip()
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # Пропускаем закрывающий ```

            # Добавляем код как моноширинный текст
            code_para = doc.add_paragraph()
            code_text = '\n'.join(code_lines)
            run = code_para.add_run(code_text)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0, 0, 128)  # Темно-синий
            continue

        # Списки (маркированные и нумерованные)
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            # Маркированный список
            list_items = []
            while i < len(lines) and (lines[i].strip().startswith('- ') or lines[i].strip().startswith('* ')):
                item_text = lines[i].strip()[2:].strip()
                list_items.append(item_text)
                i += 1

            for item in list_items:
                p = doc.add_paragraph(item, style='List Bullet')
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
            continue

        if re.match(r'^\d+\.\s', line.strip()):
            # Нумерованный список
            list_items = []
            while i < len(lines) and re.match(r'^\d+\.\s', lines[i].strip()):
                item_text = re.sub(r'^\d+\.\s', '', lines[i].strip())
                list_items.append(item_text)
                i += 1

            for item in list_items:
                p = doc.add_paragraph(item, style='List Number')
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
            continue

        # Жирный текст
        if '**' in line:
            # Заменяем **текст** на форматирование
            parts = re.split(r'\*\*(.+?)\*\*', line)
            p = doc.add_paragraph()
            for j, part in enumerate(parts):
                if j % 2 == 1:  # Нечетные части - это жирный текст
                    run = p.add_run(part)
                    run.bold = True
                else:
                    p.add_run(part)
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
            i += 1
            continue

        # Курсивный текст
        if '*' in line and '**' not in line:
            parts = re.split(r'\*(.+?)\*', line)
            p = doc.add_paragraph()
            for j, part in enumerate(parts):
                if j % 2 == 1:  # Нечетные части - это курсивный текст
                    run = p.add_run(part)
                    run.italic = True
                else:
                    p.add_run(part)
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
            i += 1
            continue

        # Обычный текст
        p = doc.add_paragraph(line)
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

        i += 1

    # Сохраняем документ
    doc.save(docx_file_path)
    print(f"Документ успешно сохранен: {docx_file_path}")


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("Использование: python md_to_word.py <input_md_file> [output_docx_file]")
        print("Пример: python md_to_word.py API_DESCRIPTION.md API_DESCRIPTION.docx")
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else input_file.replace('.md', '.docx')

    try:
        parse_markdown_to_docx(input_file, output_file)
    except FileNotFoundError:
        print(f"Ошибка: Файл '{input_file}' не найден")
        sys.exit(1)
    except Exception as e:
        print(f"Ошибка при конвертации: {e}")
        sys.exit(1)
